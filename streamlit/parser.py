from typing import List, Dict
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from table_extract import TableExtractor
from numbering import DocNumberingExtractor
class DocxParser:
    def __init__(self, docx_file):
        self.docx_file = docx_file
        self.doc = Document(docx_file)
        self.num_data = self.extract_numbering_data()
        self.content = []
        self.table_dict = TableExtractor(docx_file).extract_tables()
        _, self.x_pos_data = DocxTabPositionParser(docx_file).get_positions()

    def extract_numbering_data(self):
        try:
            num = DocNumberingExtractor(self.docx_file)
            return num.numbering_data
        except Exception as e:
            print('num data not extracted:', e)
            return {}

    def get_alignment(self, paragraph):
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: 'left',
            WD_PARAGRAPH_ALIGNMENT.CENTER: 'center',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: 'right',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'justify'
        }

        alignment = paragraph.alignment
        if alignment is None:
            return 'left'
        try:
            return alignment_map[alignment]
        except KeyError:
            if alignment == 'end':
                return 'right'
            return 'unknown'

    def add_paragraph(self, para, index, para_idx):
        font_size = para.runs[0].font.size.pt if (para.runs and para.runs[0].font.size) else 10.5
        bold_text = para.runs[0].bold if para.runs else False
        alignment = self.get_alignment(para)
        numbering = self.num_data[para_idx]['numbering'] if para_idx in self.num_data else ''
        x_pos = self.x_pos_data[para_idx] if para_idx in self.x_pos_data else 0
        if bold_text== None:
            bold_text = False
            print('None,', para.text)
        self.content.append({
            "idx": index,
            "para_idx": para_idx,
            "numbering": numbering,
            "type": "paragraph",
            "font_size": font_size,
            "x_pos": x_pos,
            "align_center": alignment == 'center',
            "is_bold": bold_text,
            "is_title": False if ((bold_text==False) and (font_size<11)) else True,            
            "text": f"{numbering} {para.text}",
            "inner_content":[]
        })
    def add_table(self, index, table_count):
        #TODO: 1줄짜리 박스면 is_ title true
        self.content.append({
            "idx": index,
            "type": "table",
            "font_size": 0,
            "x_pos":0,
            "align_center": True,
            "is_bold": False,
            "is_title": False, 
            "table_id": table_count,
            "text": f"{self.table_dict[table_count]}",
            "inner_content":[]
        })

    def parse(self):
        table_count = 0
        i = 0
        para_idx = 0
        for block in self.doc.element.body:
            if block.tag.endswith('p'):
                para = self.doc.paragraphs[para_idx]
                para._element = block
                self.add_paragraph(para, i, para_idx)
                i += 1
                para_idx += 1
            elif block.tag.endswith('tbl'):
                self.add_table(i,table_count)
                table_count += 1
                i += 1
        return self.content

    @staticmethod
    def assign_levels(content: List[Dict]) -> List[Dict]:
        font_sizes = sorted({item["font_size"] for item in content}, reverse=True)
        print(font_sizes)
        align_centers = sorted({item["align_center"] for item in content}, reverse=True)
        # is_bolds = sorted({item["is_bold"] if item["is_bold"] is not None else False for item in content}, reverse=True)
        # x_poses = sorted({item["x_pos"]  for item in content if item['font_size']== 0}, reverse=True)
        print('$$$$$$$$$$$$$$$$$$')
        # print('x_poses', x_poses)
        def _calculate_level(font_size, align_center, is_bold, x_pos):
            if is_bold ==None:
                is_bold=False
            level = 0
            if font_size in font_sizes:
                level += font_sizes.index(font_size)
            if align_center in align_centers:
                level += align_centers.index(align_center)
            if not is_bold:
                level +=1
                # level += is_bolds.index(is_bold)
            return level
        ## 1단계
        for i, item in enumerate(content):
            item["level"] = _calculate_level(item.get("font_size"), item.get("align_center"), item.get("is_bold"), item.get("x_pos"))

            # else: # table인 경우엔 이전 요소의 레벨을 가져온다
            #     item["level"] = 0 if i == 0 else content[i-1]["level"]

            item["total"] = item["level"]

        return content
    
    @staticmethod
    def get_x_pos_level(content):
        current_level = None
        current_group = []
        groups = []

        # 순회하며 그룹화
        for para in content:
            if para["is_title"] == False:
                if current_level is None:
                    current_level = para["level"]
                    current_group.append(para)
                elif para["level"] == current_level:
                    current_group.append(para)
                else:
                    if current_group:
                        groups.append((current_level, current_group))
                    current_level = para["level"]
                    current_group = [para]
            else:
                if current_group:
                    groups.append((current_level, current_group))
                current_level = None
                current_group = []

        if current_group:
            groups.append((current_level, current_group))

        # 각 그룹 내 x_pos 값을 정렬하여 x_pos_level 계산
        for level, group in groups:
            x_pos_list = [p["x_pos"] for p in group]
            unique_sorted_x_pos = sorted(set(x_pos_list))

            for para in group:
                para["x_pos_level"] = unique_sorted_x_pos.index(para["x_pos"])
                para["total"] = para["total"] + para["x_pos_level"]

        return content
    
    @staticmethod
    def put_table_level(content):
        """테이블 레벨은 직전 parargraph+1이어야 함
        """
        for i, item in enumerate(content):
            if item["type"] == "table":
                if i> 0 :
                    item["level"] = content[i-1]["total"]
                else:
                    item["level"] = 0 
                item["total"] = item["level"]

    @staticmethod
    def remove_textless_content(content: List[Dict]) -> List[Dict]:
        return [item for item in content if item['text'] != '']

    @staticmethod
    def build_tree(content: List[Dict]) -> List[Dict]:
        stack = []
        root = {"inner_content": []}


        for item in content:
            item["inner_content"] = []
            while stack and stack[-1]["level"] >= item["level"]:
                stack.pop()
            if stack:
                stack[-1]["inner_content"].append(item)
            else:
                root["inner_content"].append(item)
            stack.append(item)

        return root["inner_content"]

    @staticmethod
    def generate_html(content: List[Dict]) -> str:
        def create_html(ele):
            is_bold = "font-weight: bold;" if ele['is_bold'] else ""
            font_size = f"font-size: {ele['font_size']}px;" if ele['font_size'] > 0 else ""
            style = f"{font_size} {is_bold}"
            html = f"<div style='margin-left: {ele['total'] * 20}px; {style}'>"
            if 'inner_content' in ele:
                html += f"<span class='dropdown' onclick='toggle(this)'>{ele['text']}</span>\n"
                html += f"<div class='inner-content' style='display:none'>"
                for inner in ele['inner_content']:
                    html += create_html(inner)
                html += "</div>"
            else:
                html += f"{ele['text']}"
            html += "</div>"
            return html

        res ="".join(create_html(ele) for ele in content)

        return res
    


class DocxTabPositionParser: # docx 문서의 각 줄이 얼마나 indent되어 있는지 확인
    def __init__(self, docx_file, tab_size_in_points=36):  # 기본 탭 크기는 0.5 inch (36 points)
        self.docx_file = docx_file
        self.tab_size_in_points = tab_size_in_points
        self.doc = Document(docx_file)
        self.left_margin = self.doc.sections[0].left_margin.pt

    def get_positions(self):
        positions = []
        index_pos ={}
        for i, para in enumerate(self.doc.paragraphs):
            left_indent = para.paragraph_format.left_indent
            first_line_indent = para.paragraph_format.first_line_indent

            # 들여쓰기 계산
            left_indent = left_indent.pt if left_indent is not None else 0
            first_line_indent = first_line_indent.pt if first_line_indent is not None else 0

            # 탭 문자 수 계산
            tab_count = para.text.count('\t')
            tab_indent = tab_count * self.tab_size_in_points

            # 절대 위치는 페이지 왼쪽 여백 + 단락의 왼쪽 들여쓰기 + 첫 줄 들여쓰기 + 탭 들여쓰기
            absolute_position = self.left_margin + left_indent + first_line_indent + tab_indent
            positions.append((para.text, absolute_position))
            index_pos[i] = absolute_position

        return positions, index_pos