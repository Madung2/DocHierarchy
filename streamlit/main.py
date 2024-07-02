from docx import Document
import json
from typing import List, Dict
import streamlit as st
import streamlit.components.v1 as components
from table_extract import *
from numbering import DocNumberingExtractor

def get_alignment(paragraph):
    alignment_map = {
        0: 'left',
        1: 'center',
        2: 'right',
        3: 'justify'
    }
    return alignment_map.get(paragraph.alignment, 'unknown')

def parse_docx(file, table_dict) -> List[Dict]:
    
    document = Document(file)
    num = DocNumberingExtractor(file)
    num_data= num.numbering_data
    content = []
    table_count = 0
    i = 0

    def add_paragraph(para, index, para_idx):
        font_size = 0
        bold_text = False

        if para.runs:
            if para.runs[0].font.size:
                font_size = para.runs[0].font.size.pt

        alignment = get_alignment(para)

        if para.runs:
            if para.runs[0].bold:
                bold_text = True
        numbering = num_data[para_idx]['numbering'] if para_idx in num_data else ''
        content.append({
            "idx": index,
            "para_idx": para_idx, # this is for numbering
            "numbering": numbering,
            "type": "paragraph",
            "font_size": font_size,
            "align_center": True if alignment == 'center' else False,
            "is_bold": bold_text,
            "text": f"{numbering} {para.text}"
        })
    para_idx = 0 
    for block in document.element.body:
        
        if block.tag.endswith('p'):
            para = document.paragraphs[para_idx]
            para._element = block
            add_paragraph(para, i, para_idx)
            i += 1
            para_idx +=1
        elif block.tag.endswith('tbl'):
            content.append({
                "idx": i,
                "type": "table",
                "font_size": 0,
                "align_center": True,
                "is_bold": False,
                "table_id": table_count,
                "text": f"{table_dict[table_count]}"
            })
            table_count += 1
            i += 1

    return content
def assign_levels(content: List[Dict]) -> List[Dict]:
    # 우선 font-size 기준으로 레벨을 결정합니다.

    font_sizes = sorted({item["font_size"] for item in content}, reverse=True)
    align_centers = sorted({item["align_center"] for item in content}, reverse=True)
    is_bolds = sorted({item["is_bold"] for item in content}, reverse=True)
    def calculate_level(font_size, align_center, is_bold, font_sizes=font_sizes, align_centers=align_centers, is_bolds=is_bolds):
        level = 0
        # font_size에 따른 레벨 계산
        if font_size in font_sizes:
            level += font_sizes.index(font_size)
        # align_center에 따른 레벨 계산
        if align_center in align_centers:
            level += align_centers.index(align_center)
        # is_bold에 따른 레벨 계산
        if is_bold in is_bolds:
            level += is_bolds.index(is_bold)
        return level

    # 각 paragraph에 대해 레벨 추가
    for item in content:
        item["level"] = calculate_level(item.get("font_size"), item.get("align_center"), item.get("is_bold"))
    return content

def generate_html(content):
    def create_html(ele):
        is_bold = "font-weight: bold;" if ele['is_bold'] else ""
        font_size = f"font-size: {ele['font_size']}px;" if ele['font_size'] > 0 else ""
        style = f"{font_size} {is_bold}"

        html = f"<div style='margin-left: {ele['level'] * 20}px; {style}'>"
        if 'inner_content' in ele:
            html += f"<span class='dropdown' onclick='toggle(this)'>{ele['text']}</span>\n"
            html += f"<div class='inner-content' style='display:none'>"
            for inner in ele['inner_content']:
                html += create_html(inner)
            html += "</div>"
        else:
            html += f"{ele['text']}"
        html += "</div>"
        return html

    html = ""
    for ele in content:
        html += create_html(ele)
    return html

def remove_textless_content(content):
    processed_content = []
    for ele in content:
        if ele['text'] !='':
            
            processed_content.append(ele)
    return processed_content

def build_tree(content):
    # 트리 구조를 위한 스택
    stack = []
    root = {"inner_content": []}
    current_node = root
    
    for item in content:
        item["inner_content"] = []
        
        # 현재 노드보다 큰 레벨의 경우 스택에 추가
        while stack and stack[-1]["level"] >= item["level"]:
            stack.pop()
        
        # 스택이 비어있으면 루트에 추가
        if stack:
            stack[-1]["inner_content"].append(item)
        else:
            root["inner_content"].append(item)
        
        # 현재 노드를 스택에 추가
        stack.append(item)
    
    return root["inner_content"]



st.title("docx 문서 구조화 모듈")
st.write("----------------------")
st.write("본 프로그램은 테스트용으로 하기의 기능은 따로 모듈로 추가해야 함")
st.write("* 번호매기기 모듈")
uploaded_file = st.file_uploader("Choose a file", type=['docx', 'doc'])
st.write("----------------------")

html_code =""
if uploaded_file is not None:
    print(uploaded_file, type(uploaded_file))
    if uploaded_file.type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        content = {"error": "Invalid file type"}

    ## do table ##
    table_dict = extract_tables_from_html(uploaded_file)
    content = parse_docx(uploaded_file, table_dict)
    content = assign_levels(content)
    content = remove_textless_content(content) ## text없는 공백줄도 필요하면 이부분 주석처리 요망

    ## make html content
    html_code = generate_html(content)
    content = build_tree(content)

else:
    content ={"hello": "World!"}


tabs = ["HTML", "JSON"]

selected_tab = st.radio('Select a tab', tabs)
if selected_tab == 'HTML':
    st.markdown(html_code, unsafe_allow_html=True)
elif selected_tab == 'JSON':
    st.json(content, expanded=True)
else:
    st.markdown(html_code, unsafe_allow_html=True)



# col1, col2 = st.columns(2)
# with col1:
# with col2:
# components.html(f"{css_code}{html_code}{js_code}", height=800, scrolling=True)
